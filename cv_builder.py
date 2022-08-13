from docx import Document
from docx.shared import Inches
#import pyttsx3

#def speak(text):
    #engine = pyttsx3.init()
    #voices = engine.getProperty('voices')
    #engine.setProperty('rate', 155)
    #engine.setProperty('voice', voices[1].id)
    #engine.say(text)
    #engine.runAndWait()

document = Document()

document.add_picture('mine.jpg', width = Inches(1.3))

print(' ')
#speak("We will need your bio data")
#Personal Info
print("PERSONAL DATA")
document.add_heading('PERSONAL DETAILS')
name = input("What is your full name? Surname first ")
#speak("Hello" + name + "welcome to Sir Roll and Deborah's CV generator app")
digital_address = input("Enter your digital address? ")
email = input("Enter your email address? ")
phone = input("Enter your phone number? ")

document.add_paragraph(
    name + ' | ' + phone + ' | ' + email)

print(' ')

#Objectives
#speak("Your objectives")
print("OBJECTIVES")
document.add_heading('OBJECTIVES')
document.add_paragraph(
    input('What are you seeking for? \nType down your objectives? ')
)

print(' ')
#speak("Your missions")
#Mission Statement
print("MISSION STATEMENT")
document.add_heading('MISSION STATEMENT')
document.add_paragraph(
    input('Describe your mission ')
)

print(' ')
#speak("Your work experience")
#Work Experiences
print("WORK EXPERIENCES")
document.add_heading('WORK EXPERIENCE')
p = document.add_paragraph()

company = input('Enter the name of your previous company ')
job_title = input('Enter your job title ')
job_type = input('Enter your job type ')
print("Enter only the year")
start_date = input('Start Date ')
end_date = input('End Date ')
experience_details = input('What was your role and responsibilities at' + ' ' + company + '? ')

p.add_run(company + ' ').bold=True
p.add_run(job_title + ' ').italics=True
p.add_run(job_type + ' ').italics=True
p.add_run(start_date + '-' + end_date + ' ').italics=True
p.add_run(experience_details)

#More experience
while True:
    more_experience = input('Do you want to add more experiences? Yes or No ')
    if more_experience.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter the name of your previous company ')
        job_title = input('Enter your job title ')
        job_type = input('Enter your job type ')
        print("Enter only the year")
        start_date = input('Start Date ')
        end_date = input('End Date ')
        experience_details = input('What was your role and responsibilities at' + ' ' + company + '? ')

        p.add_run(company + ' ').bold=True
        p.add_run(job_title + ' ').italics=True
        p.add_run(job_type + ' ').italics=True
        p.add_run(start_date + '-' + end_date + ' ').italics=True
        p.add_run(experience_details)
    else:
        break

print(' ')
#speak("Your educational background")
#Education
print("EDUCATION")
document.add_heading('EDUCATION')
p = document.add_paragraph()

school = input('Enter the name of the school or university you attended ')
course = input('Enter the name of the course you offered ')
print("Enter only the year")
start_date = int(input('Start Date '))
end_date = int(input('End Date '))
certificate = input('What certificate was awarded to you at' + ' ' + school + '? ')

p.add_run(school + ' ').bold=True
p.add_run(course)
p.add_run(start_date + '-' + end_date + '\n').italics=True
p.add_run(certificate)

while True:
    another_sch = input('Do you want to add another educational background? Yes or No ')
    if another_sch.lower() == 'yes':
        p = document.add_paragraph()
        school = input('Enter the name of the school or university you attended ')
        course = input('Enter the name of the course you offered ')
        print("Enter only the year")
        start_date = int(input('Start Date '))
        end_date = int(input('End Date '))
        certificate = input('What certificate was awarded to you at' + ' ' + school + '? ')

        p.add_run(school + ' ').bold=True
        p.add_run(course)
        p.add_run(start_date + '-' + end_date + '\n').italics=True
        p.add_run(certificate)
    else:
        break

print(' ')
#speak("Your achievements")
#Achievement
print("ACHIEVEMENTS & AWARDS")
document.add_heading('ACHIEVEMENT & AWARDS')
award = input('Enter any achievement you have made or any award you have been granted ')
p=document.add_paragraph(award)
p.style = 'List Bullet'

while True:
    award = input('Do you want to add another achievement? Yes or No ')
    p=document.add_paragraph(award)
    if award.lower() == 'yes':
        award = input('Enter any achievement you have made or any award you have been granted ')
        p=document.add_paragraph(award)
        p.style = 'List Bullet'
    else:
        break

print(' ')
#speak("Your skills")
#Skills
print("SKILLS")
document.add_heading('SKILLS')
skill = input('Enter your Skills ')
p=document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    more_skills = input('Do you want to add another skills? Yes or No ')
    if more_skills.lower() == 'yes':
        skill = input('Enter your Skills ')
        p=document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

print(' ')
#speak("Your referees")
#Reference
print("REFERENCES")
document.add_heading('REFEREE')
p = document.add_paragraph()

name = input("Enter Referee's name ")
job = input("Enter Referee's job title ")
company = input("Enter Referee's company name ")
email = input("Enter Referee's email address? ")
phone = input("Enter Referee's phone number? ")

p.add_run(name + ' ').bold=True
p.add_run(job + ' ').italics=True
p.add_run(company + ' ').italics=True
p.add_run(phone + ' ').italics=True

while True:
    another_ref = input('Do you want to add anothe Referee? Yes or No ')
    if another_ref.lower() == 'yes':
        p = document.add_paragraph()
        name = input("Enter Referee's name ")
        job = input("Enter Referee's job title ")
        company = input("Enter Referee's company name ")
        email = input("Enter Referee's email address? ")
        phone = input("Enter Referee's phone number? ")

        p.add_run(name + ' ').bold=True
        p.add_run(job + ' ').italics=True
        p.add_run(company + ' ').italics=True
        p.add_run(phone + ' ').italics=True

    else:
        break


#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV GENERATOR. DEVELOPER: ASUBONTENG VINCENT"

document.save('My-CV.docx')
